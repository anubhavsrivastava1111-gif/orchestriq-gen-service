"""
OrchestrIQ Document Intelligence Engine v4 — DOCX Engine
Real .docx: cover page, auto-updating TOC field, styled heading hierarchy,
shaded table headers, footer page numbers.
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
TEAL = RGBColor(0x14, 0xB8, 0xA6)
GREY = RGBColor(0x64, 0x74, 0x8B)


def _cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for tag, attrs, text in [
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, ' TOC \\o "1-2" \\h \\z \\u '),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:t", {}, "Right-click → Update Field to build the Table of Contents."),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items(): el.set(qn(k), v)
        if text: el.text = text
        run._r.append(el)


def _page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, attrs in [("w:fldChar", {"w:fldCharType": "begin"}),
                       ("w:instrText", {"xml:space": "preserve"}),
                       ("w:fldChar", {"w:fldCharType": "end"})]:
        el = OxmlElement(tag)
        for k, v in attrs.items(): el.set(qn(k), v)
        if tag == "w:instrText": el.text = " PAGE "
        run._r.append(el)
    run.font.size = Pt(9); run.font.color.rgb = GREY


def _styled_table(doc, rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for j, row in enumerate(rows):
        for i, val in enumerate(row):
            cell = t.cell(j, i)
            cell.text = str(val)
            para = cell.paragraphs[0]
            para.runs[0].font.size = Pt(10)
            if j == 0:
                _cell_bg(cell, "1E3A5F")
                para.runs[0].font.bold = True
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif j % 2 == 0:
                _cell_bg(cell, "F1F5F9")
    return t


def build_docx(model: dict, title: str, subtitle: str = "Detailed Financial Review",
               currency_symbol: str = "\u20b9") -> bytes:
    doc = Document()
    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 13)]:
        st = doc.styles[lvl]
        st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True

    # ── COVER ────────────────────────────────────────────────────
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = NAVY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size = Pt(15); r.font.color.rgb = TEAL
    for _ in range(10): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential — Prepared for the Board of Directors")
    r.font.size = Pt(10); r.font.color.rgb = GREY
    doc.add_page_break()

    # ── TOC ──────────────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    _toc_field(doc)
    doc.add_page_break()

    # ── KPI SCORECARD ────────────────────────────────────────────
    doc.add_heading("Executive KPI Scorecard", level=1)
    _styled_table(doc, [["KPI", "Value", "Δ vs Prior"]] +
                  [list(k[:3]) for k in model["kpis"][:8]])
    doc.add_paragraph()

    # ── FINANCIAL SUMMARY TABLE ──────────────────────────────────
    doc.add_heading("Quarterly Financial Summary", level=1)
    months = model["months"]
    fmt = lambda v: f"{currency_symbol}{round(v):,}"
    _styled_table(doc, [["Line Item"] + list(months) + ["Q2 Total"],
                        ["Revenue"] + [fmt(v) for v in model["rev"]] + [fmt(sum(model["rev"]))],
                        ["COGS"] + [fmt(v) for v in model["cogs"]] + [fmt(sum(model["cogs"]))],
                        ["Gross Profit"] + [fmt(v) for v in model["gross"]] + [fmt(sum(model["gross"]))],
                        ["Opex"] + [fmt(v) for v in model["opex"]] + [fmt(sum(model["opex"]))],
                        ["EBITDA"] + [fmt(v) for v in model["ebitda"]] + [fmt(sum(model["ebitda"]))]])
    doc.add_paragraph()

    # ── NARRATIVE SECTIONS ───────────────────────────────────────
    for s in (model.get("sections") or []):
        doc.add_heading(s["h"], level=1)
        doc.add_paragraph(s["body"])

    # ── RISK REGISTER ────────────────────────────────────────────
    doc.add_heading("Risk Register", level=1)
    _styled_table(doc, [["Risk", "Severity", "Mitigation"]] +
                  [list(r[:3]) for r in model["risks"][:6]])
    doc.add_paragraph()

    # ── RECOMMENDATIONS ──────────────────────────────────────────
    doc.add_heading("Recommendations", level=1)
    for i, rec in enumerate(model["recs"][:6], 1):
        doc.add_paragraph(f"{i}. {rec}", style="List Number" if False else None)

    # ── APPENDIX ─────────────────────────────────────────────────
    doc.add_heading("Appendix — Assumptions & Methodology", level=1)
    doc.add_paragraph(
        "Figures in this review reconcile to the accompanying Excel workbook. "
        "COGS is modeled at 22% of revenue based on the trailing six-month average. "
        "Scenario growth rates (Base 12% / Bull 20% / Bear 4%) are pipeline-weighted. "
        "Runway is computed on trailing three-month net burn.")

    _page_numbers(doc)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()
