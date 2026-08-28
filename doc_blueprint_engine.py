"""
OrchestrIQ Document Intelligence Engine v4.2 — Document Blueprint Engine
Generic renderers for AI-designed PPTX / PDF / DOCX blueprints.
The AI decides the structure per request; these functions render ANY spec.

PPTX blueprint:
{"title","subtitle","slides":[
  {"type":"bullets","h":"...","kicker":"...","points":["..."],"notes":"..."},
  {"type":"chart","h":"...","chart":{"ctype":"bar|line|pie","cats":[...],"series":[["name",[nums]],...]},"notes":"..."},
  {"type":"table","h":"...","table":{"rows":[[...],...]},"notes":"..."},
  {"type":"two_col","h":"...","left":["..."],"right":["..."],"notes":"..."},
  {"type":"kpi","h":"...","kpis":[["label","value","delta"],...],"notes":"..."}]}

DOC blueprint (PDF + DOCX share it):
{"title","subtitle","sections":[
  {"h":"...","body":"paragraph","bullets":["..."],
   "table":{"rows":[[...],...]},
   "chart":{"ctype":"bar","cats":[...],"series":[["name",[nums]]]}}]}
"""
import io

# ═════════════ PPTX ═════════════
from pptx import Presentation
from pptx.util import Inches, Pt
from pptx.enum.chart import XL_CHART_TYPE
from pptx_engine import (_blank, _bar, _txt, _bullets, _notes, _header,
                         _chart, _table, NAVY, TEAL, WHITE, GREY, SW, SH)

_CT = {"bar": XL_CHART_TYPE.COLUMN_CLUSTERED, "line": XL_CHART_TYPE.LINE_MARKERS,
       "pie": XL_CHART_TYPE.PIE, "stacked": XL_CHART_TYPE.COLUMN_STACKED,
       "hbar": XL_CHART_TYPE.BAR_CLUSTERED}


def _fmt_num(v):
    """Numbers as a person writes them: 37,500 not 37500.0, 3.5 not 3.5000."""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    if f == int(f) and abs(f) < 1e15:
        return "{:,}".format(int(f))
    return "{:,.2f}".format(f)


def _chart_data_rows(cats, series, max_cols=9):
    """The chart's own numbers, as a table the reader can check.

    A chart is a claim; this is the evidence for it. Categories run across the
    top and each series becomes a row, which is the orientation that stays
    readable on a slide. If there are more categories than fit, the extra ones
    are dropped from the TABLE only - never from the chart - and the header
    says so, so nobody assumes they are seeing all of it."""
    if not cats or not series:
        return []
    show = cats[:max_cols]
    truncated = len(cats) > max_cols
    head = ["Series"] + [str(c) for c in show]
    if truncated:
        head[-1] = str(show[-1]) + " (+" + str(len(cats) - max_cols) + " more)"
    rows = [head]
    for name, vals in series[:5]:
        rows.append([_clean_cell(name, 34)] + [_fmt_num(v) for v in vals[:len(show)]])
    return rows


def _safe_chart_data(ch):
    cats = [_clean_cell(c, 24) for c in (ch.get("cats") or ["A", "B", "C"])[:12]]
    series = []
    for s in (ch.get("series") or [])[:4]:
        if isinstance(s, list) and len(s) == 2 and isinstance(s[1], list):
            vals = [float(v) if isinstance(v, (int, float)) else 0 for v in s[1][:12]]
            vals += [0] * (len(cats) - len(vals))
            # WAS str(s[0])[:30] - a blind 30-character cut. That produced the
            # legend entries you saw: "SaaS Revenue (3 contracts x $1",
            # "Audit Fee Avoidance (One-Time ", "Series A Readiness (Valuation ".
            # All severed mid-word with the bracket left open. _clean_cell cuts
            # at a WORD boundary and marks the cut, so a legend never lies about
            # what a series is.
            series.append((_clean_cell(s[0], 42), vals[:len(cats)]))
    if not series:
        series = [("Series 1", [1.0] * len(cats))]
    return cats, series


def render_pptx_blueprint(bp: dict) -> bytes:
    prs = Presentation()
    prs.slide_width = SW; prs.slide_height = SH
    title = _clean_cell(bp.get("title", "Presentation"), 88)
    subtitle = _clean_cell(bp.get("subtitle", ""), 96)

    # Title slide
    s = _blank(prs)
    _bar(s, 0, 0, SW, SH, NAVY)
    _bar(s, 0, Inches(4.6), SW, Inches(0.08), TEAL)
    _txt(s, Inches(0.9), Inches(2.6), Inches(11.5), Inches(1.4), title, 40, True, WHITE)
    if subtitle:
        _txt(s, Inches(0.9), Inches(4.9), Inches(11), Inches(0.6), subtitle, 20, False, TEAL)
    _notes(s, bp.get("opening_notes", "Open with the one-line framing of this deck."))

    # Agenda from slide headings
    heads = [_clean_cell(sl.get("h", ""), 66) for sl in (bp.get("slides") or []) if sl.get("h")]
    if len(heads) >= 3:
        s = _blank(prs); _header(s, "Agenda", "Overview")
        half = (len(heads[:12]) + 1) // 2
        _bullets(s, Inches(0.8), Inches(1.9), Inches(5.8), Inches(5), heads[:half], 16)
        _bullets(s, Inches(7.0), Inches(1.9), Inches(5.8), Inches(5), heads[half:12], 16)
        _notes(s, "Walk the agenda in 20 seconds.")

    for sl in (bp.get("slides") or [])[:24]:
        t = sl.get("type", "bullets")
        # WAS [:70] and [:40], blind cuts. They produced the headings you saw
        # with an unclosed bracket: "MONTH-END CLOSE PROCESS (GHOST STANDARD"
        # and "FINANCIAL & COMPLIANCE BENEFITS (3-YEAR".
        h = _clean_cell(sl.get("h", "Section"), 78)
        kick = _clean_cell(sl.get("kicker", ""), 52)
        s = _blank(prs); _header(s, h, kick)
        if t == "chart" and sl.get("chart"):
            cats, series = _safe_chart_data(sl["chart"])
            ctype = _CT.get(sl["chart"].get("ctype", "bar"), _CT["bar"])
            # EVERY CHART NOW SHOWS THE NUMBERS IT WAS DRAWN FROM.
            # A chart on its own cannot be checked: the reader sees a shape and
            # has to trust it. Printing the source values underneath makes the
            # figures auditable on the slide itself - if a number in the table
            # is wrong, it is visible, rather than hidden inside a picture.
            _chart(s, ctype, cats, series, Inches(0.9), Inches(1.75),
                   Inches(11.5), Inches(3.75), _clean_cell(sl["chart"].get("title", h), 70))
            drows = _chart_data_rows(cats, series)
            if drows:
                _table(s, drows, Inches(0.9), Inches(5.72), Inches(11.5), Inches(1.25))
                _txt(s, Inches(0.9), Inches(7.02), Inches(11.5), Inches(0.3),
                     "Source data for the chart above \u2014 figures as supplied by the "
                     "workspace analysis.", 9, False, GREY)
        elif t == "table" and (sl.get("table") or {}).get("rows"):
            rows = [[_clean_cell(c, 60) for c in r[:6]] for r in sl["table"]["rows"][:10]]
            ncol = len(rows[0]) if rows else 1
            rows = [(r + [""] * ncol)[:ncol] for r in rows]
            w = 11.9
            ratios = _col_ratios(rows) or [1.0 / ncol] * ncol
            _table(s, rows, Inches(0.7), Inches(1.9), Inches(w), Inches(4.6),
                   col_widths=[Inches(w * x) for x in ratios])
        elif t == "kpi" and sl.get("kpis"):
            rows = [["Metric", "Value", "Δ"]] + [[_clean_cell(x, 40) for x in (k[:3] + [""] * (3 - len(k[:3])))]
                                                for k in sl["kpis"][:8]]
            _table(s, rows, Inches(1.2), Inches(1.9), Inches(10.9), Inches(4.6))
        elif t == "two_col":
            _bullets(s, Inches(0.8), Inches(1.9), Inches(5.8), Inches(4.8),
                     [_clean_cell(p, 120) for p in (sl.get("left") or [])[:6]], 15)
            _bullets(s, Inches(7.0), Inches(1.9), Inches(5.8), Inches(4.8),
                     [_clean_cell(p, 120) for p in (sl.get("right") or [])[:6]], 15)
        else:
            # WAS str(p)[:160]. Two faults: it cut mid-word, and it left
            # markdown in place - which is why "**APPROVE**" and "**GATE 1**"
            # appeared literally on your slides with the asterisks showing.
            # _clean_cell strips the markers and cuts at a word boundary. 210
            # characters is what fits this box without overflowing it.
            _bullets(s, Inches(0.8), Inches(1.9), Inches(11.8), Inches(4.6),
                     [_clean_cell(p, 210) for p in (sl.get("points") or ["Content"])[:7]], 17)
        _notes(s, str(sl.get("notes", f"Speak to: {h}"))[:400])

    # Closing
    s = _blank(prs)
    _bar(s, 0, 0, SW, SH, NAVY)
    _txt(s, Inches(0.9), Inches(2.8), Inches(11.5), Inches(1.2), "Thank You", 44, True, WHITE)
    _txt(s, Inches(0.9), Inches(4.2), Inches(11.5), Inches(0.8),
         "Questions & Discussion", 20, False, TEAL)
    _notes(s, "Open the floor for questions.")

    # Floor: >=10 slides
    while len(prs.slides) < 10:
        s = _blank(prs); _header(s, "Supplementary", "Appendix")
        _bullets(s, Inches(0.8), Inches(1.9), Inches(11.8), Inches(4), 
                 ["Additional detail available on request"], 16)
        _notes(s, "Auto-appendix.")
    buf = io.BytesIO(); prs.save(buf)
    return buf.getvalue()


def validate_pptx_blueprint(bp) -> bool:
    return isinstance(bp, dict) and isinstance(bp.get("slides"), list) and len(bp["slides"]) >= 4


# ═════════════ PDF ═════════════
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                Spacer, Table, PageBreak, NextPageTemplate)
from pdf_engine import (_cover, _hf, _tstyle, _bar_drawing, _wrapped_table,
                        S_H1, S_BODY, S_TOC, S_SMALL, W, H)


# ── Shared table helpers ─────────────────────────────────────────────────────
# Both engines were giving EVERY column the same width regardless of content.
# In your PowerPoint that produced six columns at exactly 2.02 inches each, so
# "Typical Annual Cost per Seat" was cut off while "Player" wasted most of its
# space. Widths are now proportional to what each column actually holds.
def _col_ratios(rows):
    if not rows:
        return []
    ncol = len(rows[0])
    out = []
    for i in range(ncol):
        vals = [str(r[i]) for r in rows[1:] if i < len(r)] or [str(rows[0][i])]
        avg = sum(len(v) for v in vals) / max(len(vals), 1)
        head = len(str(rows[0][i])) if i < len(rows[0]) else 0
        # numeric columns stay narrow; long prose gets more room
        numeric = all(v.strip().replace(",", "").replace(".", "").replace("%", "")
                      .replace("-", "").replace("\u20b9", "").isdigit()
                      for v in vals if v.strip())
        w = 1.0 if numeric else max(1.0, min(3.2, avg / 14.0))
        out.append(max(w, head / 22.0))
    tot = sum(out) or ncol
    return [w / tot for w in out]
 
 
def _clean_cell(s, cap):
    """Strip markdown the renderers cannot parse, and cut at a WORD boundary."""
    t = str(s if s is not None else "")
    t = t.replace("**", "").replace("__", "").strip()
    if len(t) <= cap:
        return t
    sp = t[:cap].rfind(" ")
    return (t[:sp] if sp > cap * 0.6 else t[:cap]).rstrip(" ,;:-") + "\u2026"
 
 
def render_pdf_blueprint(bp: dict) -> bytes:
    # Session 50 made these word-safe for PPTX but NOT for PDF. This is the
    # blind [:90] that produced the cover you sent me, ending "...and Financial
    # Statement Ac" - severed mid-word, on the front page.
    title = _clean_cell(bp.get("title", "Report"), 110)
    subtitle = _clean_cell(bp.get("subtitle", "Business Report"), 130)
    buf = io.BytesIO()
    doc = BaseDocTemplate(buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
                          topMargin=2.4 * cm, bottomMargin=2 * cm)
    frame = Frame(2 * cm, 2 * cm, W - 4 * cm, H - 4.6 * cm, id="f")
    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[frame], onPage=lambda c, d: _cover(c, d, title, subtitle)),
        PageTemplate(id="body", frames=[frame], onPage=lambda c, d: _hf(c, d, title))])
    el = [NextPageTemplate("body"), PageBreak()]
    secs = (bp.get("sections") or [])[:14]
    el.append(Paragraph("Table of Contents", S_H1))
    for i, s in enumerate(secs, 1):
        # WAS [:70]. This is the contents line you saw reading "SaaS,
        # Construction, and Su".
        el.append(Paragraph(f"{i}. {_clean_cell(s.get('h','Section'), 88)}", S_TOC))
    el.append(PageBreak())
    for s in secs:
        el.append(Paragraph(_clean_cell(s.get("h", "Section"), 96), S_H1))
        if s.get("body"):
            # _clean_cell also strips ** and __, which were reaching the page.
            el.append(Paragraph(_clean_cell(s["body"], 6000), S_BODY))
        for b in (s.get("bullets") or [])[:8]:
            el.append(Paragraph("•  " + _clean_cell(b, 700), S_BODY))
        tab = (s.get("table") or {}).get("rows")
        if tab:
            # Wrap every cell (Paragraph) so long text wraps inside its column
            # instead of overflowing the right margin. Even column widths that
            # sum to exactly the frame width keep the table inside the page.
            rows = [[_clean_cell(c, 90) for c in r[:6]] for r in tab[:14]]
            # Every row padded to the header width. A short row used to shift its
            # own cells left, which is why headings sat over the wrong column.
            ncol = len(rows[0]) if rows else 1
            rows = [(r + [""] * ncol)[:ncol] for r in rows]
            avail = W - 4 * cm
            ratios = _col_ratios(rows) or [1.0 / ncol] * ncol
            t = _wrapped_table(rows, [avail * x for x in ratios])
            el.append(Spacer(1, 8)); el.append(t); el.append(Spacer(1, 8))
        ch = s.get("chart")
        if ch and ch.get("series"):
            cats, series = _safe_chart_data(ch)
            el.append(Spacer(1, 8))
            # TWO FAULTS HERE, BOTH VISIBLE IN THE PDF YOU SENT ME.
            #
            # 1. It passed series[0] ONLY. Your chart compared "Cash Basis"
            #    against "Accrual Basis" - the entire point of the exhibit - and
            #    the PDF drew one bar set. The comparison silently vanished.
            #    _bar_drawing has always handled multiple series; the caller
            #    simply never gave it more than one.
            # 2. The title was cut at 60 characters, which is why the chart read
            #    "Monthly Net Income: Cash Basis vs. Accrual Basis ($12k Annua".
            el.append(_bar_drawing(cats[:10],
                                   [(nm, vals[:10]) for nm, vals in series[:4]],
                                   _clean_cell(ch.get("title", s.get("h", "Chart")), 78)))
            el.append(Spacer(1, 6))
            # And the chart's own numbers, so the exhibit can be checked. PPTX
            # got this in Session 50; PDF is the format that most needs it,
            # because a PDF is read with nobody present to answer questions.
            drows = _chart_data_rows(cats, series)
            if drows:
                _avail = W - 4 * cm
                _r = _col_ratios(drows) or [1.0 / len(drows[0])] * len(drows[0])
                el.append(_wrapped_table(drows, [_avail * x for x in _r]))
                el.append(Paragraph("Source data for the chart above.", S_SMALL))
            el.append(Spacer(1, 8))
    el.append(Spacer(1, 12))
    el.append(Paragraph("Generated by the OrchestrIQ Document Intelligence Engine.", S_SMALL))
    doc.build(el)
    return buf.getvalue()


# ═════════════ DOCX ═════════════
from docx import Document
from docx.shared import Pt as DPt, RGBColor as DRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_engine import _toc_field, _page_numbers, _styled_table, NAVY as DNAVY, TEAL as DTEAL, GREY as DGREY


def render_docx_blueprint(bp: dict) -> bytes:
    title = _clean_cell(bp.get("title", "Document"), 110)
    subtitle = _clean_cell(bp.get("subtitle", "Business Document"), 130)
    doc = Document()
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = DPt(11)
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 13)]:
        st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = DPt(sz)
        st.font.color.rgb = DNAVY; st.font.bold = True
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = DPt(28); r.font.bold = True; r.font.color.rgb = DNAVY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size = DPt(14); r.font.color.rgb = DTEAL
    for _ in range(10): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential"); r.font.size = DPt(10); r.font.color.rgb = DGREY
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1); _toc_field(doc); doc.add_page_break()
    for s in (bp.get("sections") or [])[:16]:
        doc.add_heading(_clean_cell(s.get("h", "Section"), 96), level=1)
        if s.get("body"):
            doc.add_paragraph(_clean_cell(s["body"], 4000))
        for b in (s.get("bullets") or [])[:10]:
            doc.add_paragraph(_clean_cell(b, 400), style="List Bullet")
        tab = (s.get("table") or {}).get("rows")
        if tab:
            # WAS str(c)[:50] - a blind cut that also left ** in the cells.
            _styled_table(doc, [[_clean_cell(c, 70) for c in r[:6]] for r in tab[:15]])
            doc.add_paragraph()
        # CHARTS WERE NOT HANDLED AT ALL IN WORD. A section carrying a chart had
        # it silently dropped - the analysis appeared in the PDF and the deck,
        # and was simply absent from the Word version of the same document.
        #
        # python-docx cannot create a native Word chart, and adding a plotting
        # library to this service for one feature is not a trade worth making.
        # So Word receives the chart's DATA as a labelled table. Less pretty
        # than a picture, far better than nothing - and it is exactly the
        # verifiable source data a reader needs anyway.
        ch = s.get("chart")
        if ch and ch.get("series"):
            cats, series = _safe_chart_data(ch)
            drows = _chart_data_rows(cats, series)
            if drows:
                cap = doc.add_paragraph()
                cr = cap.add_run(_clean_cell(ch.get("title", s.get("h", "Chart")), 90))
                cr.font.bold = True; cr.font.size = DPt(10); cr.font.color.rgb = DNAVY
                _styled_table(doc, drows)
                note = doc.add_paragraph()
                nr = note.add_run("Figures as supplied by the workspace analysis.")
                nr.font.size = DPt(8); nr.font.color.rgb = DGREY
                doc.add_paragraph()
    _page_numbers(doc)
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def validate_doc_blueprint(bp) -> bool:
    return isinstance(bp, dict) and isinstance(bp.get("sections"), list) and len(bp["sections"]) >= 3
